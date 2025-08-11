from typing import List, Dict, Any, Optional, Tuple
import os
import hashlib
from pathlib import Path
import uuid
from app.embeddings.embedding_factory import BaseEmbeddingService
from app.exceptions import DocumentProcessingError, FileUploadError
from app.logger import logger


class DocumentProcessor:
    """Service for processing and chunking documents"""
    
    def __init__(
        self,
        chunk_size: int = 15000,
        chunk_overlap: int = 1000,
        embedding_service: BaseEmbeddingService = None
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.embedding_service = embedding_service
    
    async def process_document(
        self,
        file_path: str,
        file_type: str,
        meta: Optional[Dict[str, Any]] = None
    ) -> Tuple[str, List[Dict[str, Any]]]:
        """Process a document and return content and chunks"""
        
        try:
            # Extract text content based on file type
            content = await self._extract_content(file_path, file_type)
            
            # Create chunks
            chunks = await self._create_chunks(content, meta or {})
            
            logger.info(f"Processed document: {file_path}, created {len(chunks)} chunks")
            return content, chunks
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise DocumentProcessingError(f"Failed to process document: {str(e)}")
    
    async def _extract_content(self, file_path: str, file_type: str) -> str:
        """Extract text content from different file types"""
        
        file_type = file_type.lower().lstrip('.')
        
        try:
            if file_type == 'txt':
                return await self._extract_txt_content(file_path)
            elif file_type == 'pdf':
                return await self._extract_pdf_content(file_path)
            elif file_type == 'docx':
                return await self._extract_docx_content(file_path)
            elif file_type == 'md':
                return await self._extract_markdown_content(file_path)
            else:
                raise DocumentProcessingError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            raise DocumentProcessingError(f"Failed to extract content from {file_type} file: {str(e)}")
    
    async def _extract_txt_content(self, file_path: str) -> str:
        """Extract content from text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    async def _extract_pdf_content(self, file_path: str) -> str:
        """Extract content from PDF file"""
        try:
            import pypdf
            
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        content.append(f"[Page {page_num + 1}]\n{text}")
            
            return "\n\n".join(content)
            
        except ImportError:
            raise DocumentProcessingError("pypdf library not installed")
    
    async def _extract_docx_content(self, file_path: str) -> str:
        """Extract content from DOCX file"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            return "\n\n".join(content)
            
        except ImportError:
            raise DocumentProcessingError("python-docx library not installed")
    
    async def _extract_markdown_content(self, file_path: str) -> str:
        """Extract content from Markdown file"""
        # For now, treat markdown as plain text
        # In the future, you might want to parse markdown structure
        return await self._extract_txt_content(file_path)
    
    async def _create_chunks(
        self,
        content: str,
        base_meta: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Create chunks from document content"""
        
        if not content.strip():
            return []
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(content):
            # Calculate end position
            end = start + self.chunk_size
            
            # If we're not at the end, try to break at a sentence or word boundary
            if end < len(content):
                # Look for sentence ending
                sentence_break = content.rfind('.', start, end)
                if sentence_break > start:
                    end = sentence_break + 1
                else:
                    # Look for word boundary
                    word_break = content.rfind(' ', start, end)
                    if word_break > start:
                        end = word_break
            
            # Extract chunk content
            chunk_content = content[start:end].strip()
            
            if chunk_content:
                # Create chunk metadata
                chunk_meta = {
                    **base_meta,
                    "chunk_index": chunk_index,
                    "start_position": start,
                    "end_position": end,
                    "chunk_size": len(chunk_content)
                }
                
                chunks.append({
                    "id": str(uuid.uuid4()),
                    "content": chunk_content,
                    "meta": chunk_meta
                })
                
                chunk_index += 1
            
            # Move start position (with overlap)
            start = max(start + 1, end - self.chunk_overlap)
        
        return chunks
    
    async def process_chunks_with_embeddings(
        self,
        chunks: List[Dict[str, Any]],
        batch_size: int = 5000  # Process embeddings in batches to avoid token limits
    ) -> List[Dict[str, Any]]:
        """Process chunks and add embeddings in batches"""
        
        if not self.embedding_service:
            logger.warning("No embedding service provided, skipping embedding generation")
            return chunks
        
        if not chunks:
            return chunks
        
        try:
            logger.info(f"Processing embeddings for {len(chunks)} chunks in batches of {batch_size}")
            
            # Process chunks in batches to avoid token limits
            for i in range(0, len(chunks), batch_size):
                batch_end = min(i + batch_size, len(chunks))
                batch_chunks = chunks[i:batch_end]
                
                # Extract content for this batch
                batch_contents = [chunk["content"] for chunk in batch_chunks]
                
                try:
                    # Generate embeddings for this batch
                    batch_embeddings = await self.embedding_service.embed_texts(batch_contents)
                    
                    # Add embeddings to chunks in this batch
                    for j, chunk in enumerate(batch_chunks):
                        chunk["embedding"] = batch_embeddings[j]
                    
                    logger.info(f"Generated embeddings for batch {i//batch_size + 1} ({len(batch_chunks)} chunks)")
                    
                except Exception as batch_error:
                    logger.error(f"Error generating embeddings for batch {i//batch_size + 1}: {str(batch_error)}")
                    # Try processing this batch one by one if batch fails
                    for chunk in batch_chunks:
                        try:
                            embedding = await self.embedding_service.embed_text(chunk["content"])
                            chunk["embedding"] = embedding
                        except Exception as single_error:
                            logger.error(f"Failed to embed chunk {chunk['id']}: {str(single_error)}")
                            # Continue without embedding for this chunk
                            chunk["embedding"] = None
            
            # Count successfully embedded chunks
            embedded_count = sum(1 for chunk in chunks if chunk.get("embedding") is not None)
            logger.info(f"Successfully generated embeddings for {embedded_count}/{len(chunks)} chunks")
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing chunks with embeddings: {str(e)}")
            # Return chunks without embeddings rather than failing completely
            return chunks
    
    def validate_file(
        self,
        file_path: str,
        max_size: int,
        allowed_extensions: List[str]
    ) -> bool:
        """Validate uploaded file"""
        
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                raise FileUploadError("File does not exist")
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > max_size:
                raise FileUploadError(f"File size ({file_size}) exceeds maximum allowed size ({max_size})")
            
            # Check file extension
            file_extension = Path(file_path).suffix.lower()
            if file_extension not in allowed_extensions:
                raise FileUploadError(f"File type {file_extension} not allowed. Allowed types: {allowed_extensions}")
            
            return True
            
        except FileUploadError:
            raise
        except Exception as e:
            raise FileUploadError(f"File validation error: {str(e)}")
    
    def calculate_file_hash(self, file_path: str) -> str:
        """Calculate MD5 hash of file for deduplication"""
        
        hash_md5 = hashlib.md5()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            logger.error(f"Error calculating file hash: {str(e)}")
            return ""
    
    async def extract_metadata(
        self,
        file_path: str,
        file_type: str,
        additional_meta: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Extract metadata from file"""
        
        meta = additional_meta or {}
        
        try:
            # Basic file metadata
            file_stat = os.stat(file_path)
            meta.update({
                "file_path": file_path,
                "file_type": file_type,
                "file_size": file_stat.st_size,
                "file_hash": self.calculate_file_hash(file_path),
                "created_at": file_stat.st_ctime,
                "modified_at": file_stat.st_mtime
            })
            
            # Type-specific metadata
            if file_type.lower() == 'pdf':
                meta.update(await self._extract_pdf_metadata(file_path))
            elif file_type.lower() == 'docx':
                meta.update(await self._extract_docx_metadata(file_path))
            
            return meta
            
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return meta
    
    async def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF file"""
        metadata = {}
        
        try:
            import pypdf
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                metadata.update({
                    "page_count": len(pdf_reader.pages),
                    "pdf_metadata": dict(pdf_reader.metadata) if pdf_reader.metadata else {}
                })
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {str(e)}")
        
        return metadata
    
    async def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX file"""
        metadata = {}
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            core_props = doc.core_properties
            
            metadata.update({
                "docx_metadata": {
                    "author": core_props.author,
                    "title": core_props.title,
                    "subject": core_props.subject,
                    "created": core_props.created.isoformat() if core_props.created else None,
                    "modified": core_props.modified.isoformat() if core_props.modified else None,
                    "last_modified_by": core_props.last_modified_by
                }
            })
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {str(e)}")
        
        return metadata
